import os
import re
import pandas as pd
import numpy as np
import pickle

import docx
from docx import Document
from docx.enum.text import WD_COLOR_INDEX

from collections import defaultdict
from collections import Counter

import nltk
from docx.shared import Pt
from nltk.corpus import stopwords
from nltk.stem.porter import PorterStemmer
from nltk.stem import WordNetLemmatizer
from nltk.stem import SnowballStemmer as snow

from sklearn.feature_extraction.text import  TfidfVectorizer


from django.contrib import messages
from django.core.files.storage import FileSystemStorage
from django.shortcuts import render, redirect


# Create your views here


def index(request):
    global attributied, name,freq_
    if request.method == "POST":
        uploaded_file = request.FILES['document']
        if uploaded_file.name.endswith('.docx'):
            # save csv file in media folder
            savefile = FileSystemStorage()

            name = savefile.save(uploaded_file.name, uploaded_file)  # this is a name file

            # know where to save file
            d = os.getcwd()  # Current directory of the project
            file_directiry = d + '\media\\' + name
            readfile(file_directiry)
            return redirect(results)
    else:
        pass

    return render(request, 'dashboard/index.html')


##project.csv

def readfile(filename):
    ## Объявление глобальных переменных
    global model, path_document, \
        data_norm, list_text_doc, len_doc_paragraph, \
        df_docx, classes_in_doc, \
        freq_,dict_,Ypredict,dict_class_color, \
        submission,dict_class_color_otv


    ## Чтение и обрабока данных


    #data_norm = normalize_data()
    def get_model(path_model):
        with open(path_model, 'rb') as file:
            pickle_model = pickle.load(file)
        return pickle_model
    model = get_model(r'model\LinearSVC_model_2700.pkl')


    ##  Создание графика распределения по дням для каждой группы
    def load_file(path_to_file):
        doc = docx.Document(path_to_file)
        all_paragraphs = doc.paragraphs
        list_text_doc = [par.text for par in all_paragraphs]
        len_doc_paragraph = len(all_paragraphs)
        return list_text_doc, len_doc_paragraph
    list_text_doc,len_doc_paragraph = load_file('media\\' + name)

    def get_dict_class_to_color(path_color):
        dict_class_color ={}
        df = pd.read_csv(path_color)
        df.index = df['Unnamed: 0']
        f_1 = df['f1-score']
        for class_, value in zip(f_1.index, f_1.values):
            print(class_)
            if value < .25:
                dict_class_color[class_] = 'GRAY_25'
            if value < .5 and value > .25:
                dict_class_color[class_] = 'YELLOW'
            if value < .75 and value > 0.5:
                dict_class_color[class_] = 'GRAY_50'
            if value > .75:
                dict_class_color[class_] = 'TURQUOISE'

        return dict_class_color

    def reverse_dict_class_color(dict_class_color):
        for_color_class = defaultdict(list)
        for key, value in dict_class_color.items():
            if len(key) < 4:
                for_color_class[value].append(key)
            else:
                pass
        return dict(for_color_class)

    dict_class_color_otv = get_dict_class_to_color('model\\'+ 'df_class_rep.csv')

    dict_class_color = reverse_dict_class_color(dict_class_color_otv)
    print(dict_class_color)
    #list_keys_valuse_counts, list_values_valuse_counts = get_list_value_counts(data_norm)
    def pandas_df_text(text_doc):
        df = pd.DataFrame(text_doc, columns=['text'])
        return df

    df_docx = pandas_df_text(list_text_doc)
    ## Создание графика распределения сообщений по дням месяца
    def clear_text(text):
        stop_words = set(stopwords.words('russian'))
        text_cleaning_re = "[A-Za-z0-9!#$%&'()*+,./:;<=>?@[\]^_`{|}~—\"\-]+"
        tokens = []
        text = re.sub(text_cleaning_re, ' ', str(text).lower()).strip()
        text = re.sub(r'\s+', ' ', text)
        for token in text.split():
            if token not in stop_words:
                token = "".join(c for c in token if token.isalnum())
                if len(token) != 2:
                    tokens.append(token)
        return " ".join(tokens)

    def pre_process(text):
        stemmer = snow('russian')
        tokens = []
        for token in text.split():
            tokens.append(stemmer.stem(token))
        return " ".join(tokens)

    texts_ = df_docx.text.apply(lambda x: clear_text(x))
    texts_stemm = texts_.apply(lambda x: pre_process(x))

    #list_keys_mount, list_values_mount = get_list_key_value_mount(data_norm)
#Возвращает словарь в котором обозначено в каком параграфе какой класс
    def tfidf_text(texts_stemm):
        tfidf_doc = TfidfVectorizer(max_features=2745, ngram_range=(1, 4))
        X = tfidf_doc.fit_transform(texts_stemm)
        Ypredict = model.predict(X)
        Ypredict = Ypredict.tolist()
        range_ = [i for i in range(len(Ypredict))]
        dict_ = {}
        for i in range_:
            dict_[i] = Ypredict[i]
        predict_df_1  = pd.DataFrame(Ypredict, columns=['target'])
        submission = pd.concat([df_docx['text'], predict_df_1], axis=1)

        return dict_,submission,Ypredict
    dict_,submission,Ypredict = tfidf_text(texts_stemm)

    def get_info_class_loss(Ypredict):
        class_ = [str(i) for i in range(1,41)]
        class_[-1] = 'non'
        list_class = []
        for i in Ypredict:
            if i not in list_class and i != 'non':
                list_class.append(i)
        answer = sorted(list_class)
        answer.append('non')
        return answer

    def freq_class(Ypredict):
        cnt = Counter(Ypredict)
        return cnt.most_common()

    freq_=freq_class(Ypredict)

    classes_in_doc = get_info_class_loss(Ypredict)

    def get_docx(submission,dict_class_color,path):
        document = Document()
        style = document.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)
        for text,target in zip(submission.text.values.tolist(),submission.target.values.tolist()):
            text_ = '{'+target+'}'
            text_+=text
            text_+=text_
            if target != 'non':
                color = dict_class_color[target]
                if color == 'GRAY_25':
                    document.add_paragraph(
                            ).add_run(text_).font.highlight_color = WD_COLOR_INDEX.GRAY_25
                if color == 'YELLOW':
                    document.add_paragraph(
                            ).add_run(text_).font.highlight_color = WD_COLOR_INDEX.YELLOW
                if color == 'GRAY_50':
                    document.add_paragraph(
                            ).add_run(text_).font.highlight_color = WD_COLOR_INDEX.GRAY_50
                if color == 'TURQUOISE':
                    document.add_paragraph(
                            ).add_run(text_).font.highlight_color = WD_COLOR_INDEX.TURQUOISE
            else:
                document.add_paragraph().add_run(text_)
        document.save(path)
    path_to_doc = 'answer\\'+'predict_model'+name
    print(path_to_doc)
    get_docx(submission, dict_class_color_otv, path_to_doc)
def results(request):
    context = {
        'freq_': freq_,
        'classes_in_doc': classes_in_doc,
        'dict_class_color': dict_class_color
    }

    return render(request, 'dashboard/results.html', context)

def nlp_results(request):
    list_path_photo = []
    return render(request,'dashboard/nlp_results.html')